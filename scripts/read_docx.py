"""Extract text from Word documents in project documents folder."""

import sys

from docx import Document


def read_docx(filepath):
    """Read and print all text from a Word document."""
    try:
        doc = Document(filepath)
        print(f"\n{'='*80}")
        print(f"DOCUMENT: {filepath}")
        print(f"{'='*80}\n")

        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)

        # Also check tables
        for table in doc.tables:
            print("\n[TABLE]")
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    print(row_text)

        print(f"\n{'='*80}\n")

    except Exception as e:
        print(f"Error reading {filepath}: {e}")


if __name__ == "__main__":
    # Read both documents
    read_docx("project documents/Dissertation_Checklist.docx")
    read_docx("project documents/COMPREHENSIVE BLUEPRINT.docx")
